"""Document processing and chunking for Local RAG (한글 최적화)"""

import hashlib
import json
import re
import unicodedata
from pathlib import Path
from dataclasses import dataclass
from typing import List, Optional, Generator
from tqdm import tqdm


@dataclass
class Document:
    """Represents a document chunk"""
    id: str
    content: str
    source: str
    chunk_index: int
    metadata: dict


class KoreanTextProcessor:
    """한글 텍스트 전처리 유틸리티 (정확도 향상 버전)"""

    # 한글 유니코드 범위
    HANGUL_PATTERN = re.compile(r'[가-힣ㄱ-ㅎㅏ-ㅣ]')

    # 한국어 문장 종결 패턴 (더 정밀한 패턴)
    KOREAN_SENTENCE_ENDINGS = re.compile(
        r'(?<=[다요죠])\.\s+|'  # 다., 요., 죠. 뒤
        r'(?<=습니다)\.\s+|'    # 습니다. 뒤
        r'(?<=했다)\.\s+|'      # 했다. 뒤
        r'(?<=된다)\.\s+|'      # 된다. 뒤
        r'(?<=이다)\.\s+|'      # 이다. 뒤
        r'(?<=[.!?])\s+(?=[A-Z가-힣])'  # 문장부호 뒤 대문자/한글
    )

    # 불용어 리스트 (검색 정확도 향상용)
    STOPWORDS = {
        '그', '저', '것', '수', '등', '및', '또', '또는', '그리고', '하지만', '그러나',
        '이', '그것', '저것', '무엇', '어떤', '모든', '각', '매', '위', '아래',
        '있다', '없다', '되다', '하다', '이다', '같다', '않다',
        'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
        'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
        'this', 'that', 'these', 'those', 'it', 'its', 'and', 'or', 'but', 'if', 'then'
    }

    @staticmethod
    def normalize(text: str) -> str:
        """한글 정규화 (자모 분리된 텍스트 결합 등)"""
        # NFC 정규화로 자모 결합
        text = unicodedata.normalize('NFC', text)
        # 연속 공백 제거
        text = re.sub(r'\s+', ' ', text)
        # 특수문자 정리 (의미있는 것만 유지)
        text = re.sub(r'[^\w\s가-힣ㄱ-ㅎㅏ-ㅣ.,!?:;\-()\'\"@#$%&*+=/]', ' ', text)
        return text.strip()

    @staticmethod
    def is_korean(text: str) -> bool:
        """텍스트가 한글을 포함하는지 확인"""
        return bool(KoreanTextProcessor.HANGUL_PATTERN.search(text))

    @staticmethod
    def get_korean_ratio(text: str) -> float:
        """텍스트 내 한글 비율 계산"""
        if not text:
            return 0.0
        korean_chars = len(KoreanTextProcessor.HANGUL_PATTERN.findall(text))
        return korean_chars / len(text.replace(' ', ''))

    @staticmethod
    def split_sentences(text: str) -> List[str]:
        """한글/영어 혼합 문장 분리 (개선된 버전)"""
        # 정규식으로 문장 분리
        sentences = KoreanTextProcessor.KOREAN_SENTENCE_ENDINGS.split(text)

        result = []
        for sent in sentences:
            sent = sent.strip()
            if sent and len(sent) > 5:  # 너무 짧은 문장 제외
                result.append(sent)

        return result if result else [text.strip()] if text.strip() else []

    @staticmethod
    def tokenize_simple(text: str) -> List[str]:
        """간단한 한글 토큰화 (공백 + 조사 분리)"""
        # 기본 공백 분리
        tokens = text.split()

        # 조사 패턴 (확장된 패턴)
        josa_pattern = re.compile(
            r'^(.+?)(은|는|이|가|을|를|에|에서|로|으로|와|과|의|도|만|까지|부터|'
            r'에게|한테|께|처럼|같이|보다|라고|라는|이라고|이라는|으로서|로서|'
            r'으로써|로써|이며|며|이고|고|이나|나|이란|란|이든|든)$'
        )

        expanded = []
        for token in tokens:
            # 한글이 포함된 경우만 조사 분리
            if KoreanTextProcessor.HANGUL_PATTERN.search(token):
                match = josa_pattern.match(token)
                if match and len(match.group(1)) > 0:
                    stem = match.group(1)
                    # 불용어가 아닌 경우만 추가
                    if stem.lower() not in KoreanTextProcessor.STOPWORDS:
                        expanded.append(stem)
                elif token.lower() not in KoreanTextProcessor.STOPWORDS:
                    expanded.append(token)
            elif token.lower() not in KoreanTextProcessor.STOPWORDS:
                expanded.append(token)

        return expanded

    @staticmethod
    def extract_keywords(text: str, max_keywords: int = 10) -> List[str]:
        """핵심 키워드 추출"""
        tokens = KoreanTextProcessor.tokenize_simple(text)

        # 빈도수 계산
        from collections import Counter
        word_freq = Counter(tokens)

        # 빈도 높은 키워드 반환 (2글자 이상)
        keywords = [word for word, _ in word_freq.most_common(max_keywords * 2)
                   if len(word) >= 2]

        return keywords[:max_keywords]


class DocumentProcessor:
    """Process and chunk documents for RAG (한글 최적화)"""

    FINGERPRINT_VERSION = "source-sync-v2"

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50,
                 use_korean_normalization: bool = True):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_korean_normalization = use_korean_normalization
        self.korean_processor = KoreanTextProcessor()

    def _generate_id(self, content: str, source: str, chunk_index: int) -> str:
        """Generate unique ID for a chunk"""
        hash_input = f"{source}:{chunk_index}:{content[:100]}"
        return hashlib.md5(hash_input.encode()).hexdigest()[:16]

    def _build_source_metadata(self, text: str) -> dict:
        """Build metadata shared by all chunks from the same source."""
        normalized_text = KoreanTextProcessor.normalize(text) if self.use_korean_normalization else re.sub(r'\s+', ' ', text).strip()
        fingerprint_payload = json.dumps(
            {
                "version": self.FINGERPRINT_VERSION,
                "text": normalized_text,
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap,
                "use_korean_normalization": self.use_korean_normalization,
            },
            ensure_ascii=True,
            sort_keys=True,
        )
        return {
            "source_fingerprint": hashlib.sha1(fingerprint_payload.encode("utf-8")).hexdigest()[:16],
            "source_char_count": len(normalized_text),
            "fingerprint_version": self.FINGERPRINT_VERSION,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
        }

    def _split_into_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks (한글 문장 경계 인식)"""
        # 한글 정규화
        if self.use_korean_normalization:
            text = KoreanTextProcessor.normalize(text)
        else:
            text = re.sub(r'\s+', ' ', text).strip()

        if len(text) <= self.chunk_size:
            return [text] if text else []

        chunks = []
        start = 0
        is_korean = KoreanTextProcessor.is_korean(text)

        while start < len(text):
            end = start + self.chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # 한글 문서인 경우 한국어 문장 종결 패턴 우선 사용
                if is_korean:
                    # 한국어 종결어미 패턴
                    korean_endings = ['다. ', '요. ', '죠. ', '습니다. ', '했다. ', '된다. ',
                                     '이다. ', '. ', '! ', '? ', '。', '\n']
                    best_end = -1
                    for sep in korean_endings:
                        last_sep = text.rfind(sep, start, end)
                        if last_sep > start + self.chunk_size // 3:
                            if last_sep > best_end:
                                best_end = last_sep + len(sep)
                    if best_end > 0:
                        end = best_end
                else:
                    # 영어 문장 종결
                    for sep in ['. ', '! ', '? ', '\n']:
                        last_sep = text.rfind(sep, start, end)
                        if last_sep > start + self.chunk_size // 2:
                            end = last_sep + 1
                            break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - self.chunk_overlap if end < len(text) else end

        return chunks

    def process_text(self, text: str, source: str = "text") -> List[Document]:
        """Process raw text into document chunks"""
        chunks = self._split_into_chunks(text)
        source_metadata = self._build_source_metadata(text)

        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                id=self._generate_id(chunk, source, i),
                content=chunk,
                source=source,
                chunk_index=i,
                metadata={
                    **source_metadata,
                    "total_chunks": len(chunks),
                    "is_korean": KoreanTextProcessor.is_korean(chunk)
                }
            )
            documents.append(doc)

        return documents

    def process_markdown(self, filepath: Path) -> List[Document]:
        """Process markdown file"""
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        source_metadata = self._build_source_metadata(content)

        # Split by headers for better context preservation, including headings at the start of the file.
        sections = re.split(r'(?m)^(#{1,6}\s+.+)$', content)

        documents = []
        current_header = ""
        chunk_idx = 0

        for i, section in enumerate(sections):
            if re.match(r'^#{1,6}\s+', section):
                current_header = section.strip()
            elif section.strip():
                # Add header to section for context
                text = f"{current_header}\n{section}" if current_header else section
                chunks = self._split_into_chunks(text)

                for chunk in chunks:
                    doc = Document(
                        id=self._generate_id(chunk, str(filepath), chunk_idx),
                        content=chunk,
                        source=str(filepath),
                        chunk_index=chunk_idx,
                        metadata={
                            **source_metadata,
                            "header": current_header,
                            "type": "markdown",
                        }
                    )
                    documents.append(doc)
                    chunk_idx += 1

        return documents

    def process_pdf(self, filepath: Path) -> List[Document]:
        """Process PDF file"""
        try:
            from pypdf import PdfReader
        except ImportError:
            print("pypdf not installed. Install with: pip install pypdf")
            return []

        reader = PdfReader(filepath)
        all_text = ""

        for page in reader.pages:
            all_text += page.extract_text() + "\n"

        return self.process_text(all_text, str(filepath))

    def process_docx(self, filepath: Path) -> List[Document]:
        """Process Word document"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return []

        doc = DocxDocument(filepath)
        all_text = "\n".join([para.text for para in doc.paragraphs])

        return self.process_text(all_text, str(filepath))

    def process_file(self, filepath: Path) -> List[Document]:
        """Process a file based on its extension"""
        filepath = Path(filepath)

        if not filepath.exists():
            print(f"File not found: {filepath}")
            return []

        ext = filepath.suffix.lower()

        if ext == '.md':
            return self.process_markdown(filepath)
        elif ext == '.pdf':
            return self.process_pdf(filepath)
        elif ext == '.docx':
            return self.process_docx(filepath)
        elif ext in ['.txt', '.py', '.js', '.ts', '.json', '.yaml', '.yml']:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            return self.process_text(content, str(filepath))
        else:
            print(f"Unsupported file type: {ext}")
            return []

    def process_directory(self, directory: Path, extensions: Optional[List[str]] = None) -> Generator[Document, None, None]:
        """Process all files in a directory"""
        directory = Path(directory)

        if extensions is None:
            extensions = ['.txt', '.md', '.pdf', '.docx', '.py', '.js', '.ts', '.json', '.yaml', '.yml']

        files = []
        for ext in extensions:
            files.extend(directory.rglob(f"*{ext}"))

        for filepath in tqdm(files, desc="Processing files"):
            documents = self.process_file(filepath)
            for doc in documents:
                yield doc
